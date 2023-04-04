from docx import Document


def checkAndChange(document, replaceDict):
    """
    遍历word中的所有 paragraphs，在每一段中发现含有key 的内容，就替换为 value 。
   （key 和 value 都是replace_dict中的键值对。）
    """
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            print(para.runs[i].text)
            for key, value in replaceDict.items():
                if key in para.runs[i].text:
                    print(key + "->" + value)
                    para.runs[i].text = para.runs[i].text.replace(key, value)
    return document


def checkAndChangeTable(document, replaceDict):
    tableNum = len(document.tables)
    for i in range(0, tableNum):
        rowNum = len(document.tables[i].rows)
        for j in range(0, rowNum):
            colNum = len(document.tables[i].rows[j].cells)
            for k in range(0, colNum):
                cell = document.tables[i].rows[j].cells[k]
                # print("cell: " + cell)
                for key, value in replaceDict.items():
                    if key == cell.text:
                        print(key + "->" + value)
                        cell.text = cell.text.replace(key, value)


def generateWordDocuments(tempFile,replaceDict,rowData,colExternal):
    doc = Document(tempFile)
    newDict = {}
    for key in replaceDict.keys():
        newDict[key] = rowData[replaceDict[key]]
    print(newDict)
    checkAndChangeTable(doc,newDict)
    doc.save(rowData[colExternal]+".docx")

